# -*- coding: utf-8 -*-
"""
Created on Mon Jan 30 11:41:29 2023

@author: enado
"""

import docx

# Crea un nuevo documento de Word
doc = docx.Document()

# Abre el archivo de texto
with open("Bayern München.txt", "r") as file:
    text = file.read()

# Separa el texto en 16 campos
fields = text.split("\n")[:16]

# Agrega cada campo como un párrafo en el documento de Word
for field in fields:
    doc.add_paragraph("Jugador")
    doc.add_paragraph("Nacimiento")
    doc.add_paragraph("Edad")
    doc.add_paragraph("Nación")
    doc.add_paragraph("Altura")
    doc.add_paragraph("Peso")
    doc.add_paragraph("PJ")
    doc.add_paragraph("Goles")
    doc.add_paragraph("Asistencias")
    doc.add_paragraph("TA")
    doc.add_paragraph("TAR")
    doc.add_paragraph("TR")
    doc.add_paragraph("Desde")
    doc.add_paragraph("De")
    doc.add_paragraph("BL")
    doc.add_paragraph("Número")

# Guarda el documento de Word
doc.save("Bayern.docx")
